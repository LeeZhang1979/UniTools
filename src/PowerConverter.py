import xlwings as xw
import wx
from docx import Document
from docx.shared import Pt

FIRST_COL_LABEL = 45
FIRST_COL_TEXT = 115

SECOND_COL_LABEL = 230
SECOND_COL_TEXT = 320

THIRD_COL_LABEL = 430
THIRD_COL_TEXT = 510

FOURTH_COL_LABEL = 620
FOURTH_COL_TEXT = 710

LABEL_WIDTH = 70
TEXT_WIDTH = 90

LABEL_TEXT_HEIGHT = 20
FIRST_ELE_X = 10
FIRST_ELE_Y = 10

LINE_HEIGHT = 3
LINE_WIDTH = 840

doc2Header = ['RZJS-B8']
doc2Tables = ['RZJS-B8', 'RZJS-D10', 'RZJS-D14', 'RZJS-C8', 'RZJS-F10', 'RZJS-G10', 'RZJS-H10', 'RZJS-H20', 'RZJS-H23',
              'RZJS-O17', 'RZJS-C20', 'RZJS-E20', 'RZJS-M10', 'RZJS-M14', 'RZJS-N10', 'RZJS-G8', ]

# 打开docx文件
doc = Document('./res/计算报告-模板.docx')
RZJSString = ['RZJS-C3', 'RZJS-D10', 'RZJS-D14', 'RZJS-B5', 'RZJS-C5', 'RZJS-C8', 'RZJS-F10', 'RZJS-G10', 'RZJS-N10',
              'RZJS-N11', 'RZJS-M10', 'RZJS-M14', 'RZJS-M15', 'RZJS-H20', 'RZJS-H23', 'RZJS-G8', 'RZJS-C20', 'RZJS-C23',
              'RZJS-E20', 'RZJS-E23', 'RZJS-E28', 'RZJS-B17', 'RZJS-B28', 'RZJS-I17', 'RZJS-C21', 'RZJS-B8',
              'RZJS-G17', 'RZJS-K17', 'RZJS-J17', 'RZJS-J21', 'RZJS-H21', 'RZJS-J47', 'RZJS-J80', 'RZJS-H17',
              'RZJS-B31', 'RZJS-K219', 'RZJS-K220', 'RZJS-K221', 'RZJS-K222', 'RZJS-K228', 'RZJS-B32', 'RZJS-D5',
              'RZJS-B191', 'RZJS-B33', 'RZJS-B192', 'RZJS-B36', 'RZJS-B193', 'RZJS-C193', 'RZJS-B38', 'RZJS-B194',
              'RZJS-B195', 'RZJS-G38', 'RZJS-H38', 'RZJS-B196', 'RZJS-G39', 'RZJS-H39', 'RZJS-B197', 'RZJS-C197',
              'RZJS-G45', 'RZJS-B44', 'RZJS-D44', 'RZJS-B198', 'RZJS-B212', 'RZJS-D46', 'RZJS-B213', 'RZJS-B47',
              'RZJS-B214', 'RZJS-M53', 'RZJS-I194',
              'RZJS-I192', 'RZJS-I193', 'RZJS-I191', 'RZJS-I195', 'RZJS-B64', 'RZJS-C64', 'RZJS-D64', 'RZJS-B65',
              'RZJS-C65', 'RZJS-D65', 'RZJS-B66', 'RZJS-C66', 'RZJS-D66', 'RZJS-B67', 'RZJS-C67', 'RZJS-D67',
              'RZJS-B69', 'RZJS-C69', 'RZJS-D69', 'RZJS-B71', 'RZJS-B75', 'RZJS-D75', 'RZJS-B73', 'RZJS-B70',
              'RZJS-L80', 'RZJS-L81', 'RZJS-L86', 'RZJS-C92', 'RZJS-D79']
QTJSString = ['QTJS-T8', 'QTJS-E22', 'QTJS-M23', 'QTJS-E26', 'QTJS-E27', 'QTJS-F22', 'QTJS-F23', 'QTJS-F26', 'QTJS-F27',
              'QTJS-E32', 'QTJS-E33', 'QTJS-I35', 'QTJS-I43', 'QTJS-E32', 'QTJS-E33', 'QTJS-J35', 'QTJS-J43',
              'QTJS-C52', 'QTJS-F32', 'QTJS-F33', 'QTJS-E36', 'QTJS-F36', 'QTJS-E37', 'QTJS-F37', 'QTJS-E38', 'QTJS-F38',
              'QTJS-C60', 'QTJS-G54', 'QTJS-C55', 'QTJS-C54', 'QTJS-G65', 'QTJS-C61', 'QTJS-G56', 'QTJS-C67',
              'QTJS-C66', 'QTJS-C57', 'QTJS-C62', 'QTJS-C58', 'QTJS-G58', 'QTJS-G61', 'QTJS-C65', 'QTJS-G59',
              'QTJS-J61', 'QTJS-C2', 'QTJS-C3', 'QTJS-C4', 'QTJS-G52', 'QTJS-G53', 'QTJS-L36', 'QTJS-T8']
baiFenHao = ['H20', 'H23', 'J21', 'H21', 'G10', 'E23', 'C23', ]
doc2BaiFenHao = ['H20']

rLiangWeiXiaoShu = ['J80', 'E32', 'K219', 'K228', 'I17', 'B193','M53', 'C83']

qLiangWeiXiaoShuR = ['E27', 'F27', 'E33', 'F33', 'K228']

wuWeiXiaoShu = ['B50']

filePathShuRu = './res/短路力计算表格.xls'
sheetShuRu = '输入表'
sheetShuChu = '输出表'
wToDuanLuLi = [('J10', 'B8'), ('p10','b8'), ('j11', 'd10'), ('p11', 'd14'), ('j15', 'h24'), ('p15', 'h24'), ('j39', 'c69'),
               ('p39', 'b37'), ('j40','f86'), ('p40', 'e53'), ('j41', 'e201'), ('p41', 'd194'), ('j43', 'd91'), ('p43', 'e58'),
               ('j44', 'f76'), ('p44', 'f43'), ('j45', 'g76'), ('p45', 'g43'), ('j46', 'g78'), ('p46', 'g45'), ('j47', 'e79'),
               ('p47', 'e46'),  ('j48', 'b75'), ('p48', 'b42'), ('j49', 'e75'), ('p49', 'e42'), ('j57', 'c190')]

sheetZhuCai = '主材'
filePathChengBen = './res/应该成本模板 变压器-35kV-20210810.V2.0(1).xlsx'
zhuCaiQT = [('e3', 'c52'), ('e7', 'c64'), ('e8', 'd58')]
zhuCaiRZ = [('e4', 'c85'), ('e5', 'c52')]

shuChuString = ['输出表-G5', '输出表-O5', '输出表-G6', '输出表-O6', '输出表-G7', '输出表-O7', '输出表-G8', '输出表-O8',
                '输出表-G9', '输出表-O9', '输出表-K10', '输出表-G12', '输出表-O12', '输出表-G15', '输出表-O15', '输出表-G16',
                '输出表-O16', '输出表-Q26', '输出表-V26', '输出表-AB26', '输出表-Q27', '输出表-V27', '输出表-AB27', '输出表-Q28',
                '输出表-Q30', '输出表-V30', '输出表-Q33', '输出表-V33', '输出表-AB33', '输出表-Q34', '输出表-V34', '输出表-AB34',
                '输出表-Q36', '输出表-V36', '输出表-N37', '输出表-V28', '输出表-AB28', '输出表-N31']

sheet1Name = "RZJS"
sheet2Name = "QTJS"
# 打开excel文件
appEx = xw.App(visible=False, add_book=False)
appEx.screen_updating = False
filePath = "./res/35kV计算单-全铜5_6.xls"
wb = appEx.books.open(filePath)
sheet = wb.sheets[sheet1Name]

tempTest = ['RZJS-C51', 'RZJS-C52', 'RZJS-B48', 'RZJS-B49', 'RZJS-B50', 'RZJS-C82', 'RZJS-C83', 'RZJS-C84', 'RZJS-C85']

def ExportDoc(event):
    dlg = wx.MessageDialog(None, u"开始导出，请稍等两分钟，将保存在程序目录", u"温馨提示", wx.YES_DEFAULT)
    if dlg.ShowModal() == wx.ID_YES:
        dlg.Destroy()
    sheet1 = wb.sheets[sheet1Name]
    # 成本打开和写入并保存

    chengBen = appEx.books.open(filePathChengBen)
    sheetChengBen = chengBen.sheets[sheetZhuCai]
    for tup in zhuCaiRZ:
        text = sheet1.range(tup[1]).value
        sheetChengBen.range(tup[0]).value = text
    sheet1 = wb.sheets[sheet2Name]
    for tup in zhuCaiQT:
        text = sheet1.range(tup[1]).value
        sheetChengBen.range(tup[0]).value = text
    chengBen.save("./应该成本模板 变压器-35kV-20210810.V2.0.xls")
    chengBen.close()

    sheet1 = wb.sheets[sheet1Name]
    # 短路力计算
    shuRuJiSuan = appEx.books.open(filePathShuRu)
    sheetJiSuanLi = shuRuJiSuan.sheets[sheetShuRu]
    for tup in wToDuanLuLi:
        sheetJiSuanLi.range(tup[0]).value = sheet1.range(tup[1]).value

    sheetJiSuanLi = shuRuJiSuan.sheets[sheetShuChu]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for strc1 in shuChuString:
                    if strc1 in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                cell1Name = strc1.split('-')[1]
                                replaceStr = str(sheetJiSuanLi.range(cell1Name).value)
                                # if cell1Name in rLiangWeiXiaoShu:
                                #     replaceStr = str(round(sheet1.range(cell1Name).value, 2))
                                # if cell1Name in wuWeiXiaoShu:
                                #     replaceStr = str(round(sheet1.range(cell1Name).value, 6))
                                # if cell1Name in baiFenHao:
                                #     replaceStr = str(round(sheet1.range(cell1Name).value * 100, 2))
                                cell.text = cell.text.replace(strc1, replaceStr)
                                replaceStr = ''
    shuRuJiSuan.save()
    shuRuJiSuan.close()

    # appEx.quit()  #退出app

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                sheet1 = wb.sheets[sheet1Name]
                for strc1 in tempTest:
                    if strc1 in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                cell1Name = strc1.split('-')[1]
                                replaceStr = str(sheet1.range(cell1Name).value)
                                if cell1Name in rLiangWeiXiaoShu:
                                    replaceStr = str(round(sheet1.range(cell1Name).value, 2))
                                if cell1Name in wuWeiXiaoShu:
                                    replaceStr = str(round(sheet1.range(cell1Name).value, 6))
                                if cell1Name in baiFenHao:
                                    replaceStr = str(round(sheet1.range(cell1Name).value * 100, 2))
                                cell.text = cell.text.replace(strc1, replaceStr)
                                replaceStr = ''


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                sheet1 = wb.sheets[sheet1Name]
                for strc1 in RZJSString:
                    if strc1 in cell.text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                cell1Name = strc1.split('-')[1]
                                replaceStr = str(sheet1.range(cell1Name).value)
                                if cell1Name in rLiangWeiXiaoShu:
                                    replaceStr = str(round(sheet1.range(cell1Name).value, 2))
                                if cell1Name in wuWeiXiaoShu:
                                    replaceStr = str(round(sheet1.range(cell1Name).value, 6))
                                if cell1Name in baiFenHao:
                                    replaceStr = str(round(sheet1.range(cell1Name).value * 100, 2))
                                cell.text = cell.text.replace(strc1, replaceStr)
                                replaceStr = ''


    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                sheet1 = wb.sheets[sheet1Name]
                for strc in QTJSString:
                    if strc in cell.text:
                        cell2Name = strc.split('-')[1]
                        sheet2 = wb.sheets[sheet2Name]
                        replaceStrq = str(sheet2.range(cell2Name).value)
                        if cell2Name in qLiangWeiXiaoShuR:
                            replaceStrq = str(round(sheet2.range(cell2Name).value, 2))
                        if cell2Name in baiFenHao:
                            replaceStrq = str(round(sheet2.range(cell2Name).value * 100, 2))
                        cell.text = cell.text.replace(strc, replaceStrq)
                        replaceStrq = ''

    for str11 in RZJSString:
        cell1Name = str11.split('-')[1]
        replaceStr = str(sheet1.range(cell1Name).value)
        if cell1Name in rLiangWeiXiaoShu:
            replaceStr = str(round(sheet1.range(cell1Name).value, 2))
        if cell1Name in baiFenHao:
            replaceStr = str(round(sheet1.range(cell1Name).value * 100, 2))
        for para in doc.paragraphs:
            if str11 in para.text:
                para.text = para.text.replace(str11, replaceStr)

    for str2 in QTJSString:
        cell2Name = str2.split('-')[1]
        replaceStrq = str(sheet2.range(cell2Name).value)
        if cell2Name in qLiangWeiXiaoShuR:
            replaceStrq = str(round(sheet2.range(cell2Name).value, 2))
        if cell2Name in baiFenHao:
            replaceStrq = str(round(sheet2.range(cell2Name).value * 100, 2))
        for para in doc.paragraphs:
            if str2 in para.text:
                para.text = para.text.replace(str2, replaceStrq)

    for name in RZJSString:
        for section in doc.sections:
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if name in cell.text:
                            sheet1 = wb.sheets[sheet1Name]
                            cell1Name = name.split('-')[1]
                            replaceStr = str(sheet1.range(cell1Name).value)
                            cell.text = cell.text.replace(name, replaceStr)

    doc.styles['Normal'].font.size = Pt(9)
    doc.save("./计算报告-结果.docx")

    doc2 = Document('./res/PA-21-XXX V5风力发电机组35kV 8800kVA油浸式变压器技术协议_供应商-模板.docx')
    for table2 in doc2.tables:
        for row in table2.rows:
            for cell in row.cells:
                for strdoc2 in doc2Tables:
                    if strdoc2 in cell.text:
                        sheet1 = wb.sheets[sheet1Name]
                        cell1Name = strdoc2.split('-')[1]
                        replaceStr = str(sheet1.range(cell1Name).value)
                        if cell1Name in doc2BaiFenHao:
                            replaceStr = str(round(sheet1.range(cell1Name).value * 100, 2))
                        cell.text = cell.text.replace(strdoc2, replaceStr)
    for str2d in doc2Header:
        cell1Name = str2d.split('-')[1]
        replaceStr = str(sheet1.range(cell1Name).value)
        for para in doc2.paragraphs:
            if str2d in para.text:
                para.text = para.text.replace(str2d, replaceStr)

    for str2d in doc2Header:
        for section in doc2.sections:
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if str2d in cell.text:
                            sheet1 = wb.sheets[sheet1Name]
                            cell1Name = str2d.split('-')[1]
                            replaceStr = str(sheet1.range(cell1Name).value)
                            cell.text = cell.text.replace(str2d, replaceStr)
    doc2.save("./PA-21-XXX V5风力发电机组35kV 8800kVA油浸式变压器技术协议_供应商-结果.docx")
    dlg = wx.MessageDialog(None, u"已保存在程序目录", u"导出完成", wx.YES_DEFAULT)
    if dlg.ShowModal() == wx.ID_YES:
        dlg.Destroy()



#def onExit():
#    appEx.quit()
app = wx.App()
frame = wx.Frame(None, title="变压器计算程序", pos=(400, 100), size=(840, 520))
#frame.Bind(wx.EVT_CLOSE,onExit())
nb = wx.Notebook(frame)
p1 = MyPanel1(nb)
p2 = MyPanel2(nb)

nb.AddPage(p1, "电气参数")
nb.AddPage(p2, "绕组计算")
frame.Show()
app.MainLoop()

class MyPanel1(wx.Panel):
    def __init__(self, parent):
        super(MyPanel1, self).__init__(parent)
        # 第一行
        basicParams = wx.StaticText(self, label='基\n\n本\n\n参\n\n数', pos=(FIRST_ELE_X, FIRST_ELE_Y),
                                    size=(LABEL_TEXT_HEIGHT, 100))
        basicParams.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        basicParams.SetForegroundColour('red')
        m_staticText120 = wx.StaticText(self, wx.ID_ANY, u"产品型号", pos=(FIRST_COL_LABEL,LABEL_TEXT_HEIGHT),
                                        size=(70, 20))
        self.m_textCtrlC3 = wx.TextCtrl(self, wx.ID_ANY, u"SSP-8800/35", pos=(FIRST_COL_TEXT, LABEL_TEXT_HEIGHT),
                                        size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlC3.Bind(wx.EVT_TEXT_ENTER, self.EndTextC3)

        m_staticText122 = wx.StaticText(self, wx.ID_ANY, u"额定容量kVA", pos=(SECOND_COL_LABEL, LABEL_TEXT_HEIGHT),
                                        size=(90, 20), style=wx.ALIGN_CENTRE_VERTICAL)
       # m_staticText122.SetFont(wx.Font(9, wx.DEFAULT, wx.NORMAL, wx.NORMAL))
        self.m_textCtrlB8 = wx.TextCtrl(self, wx.ID_ANY, u"8800", pos=(SECOND_COL_TEXT, LABEL_TEXT_HEIGHT), size=(TEXT_WIDTH, 20),
                                        style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlB8.Bind(wx.EVT_TEXT_ENTER, self.EndTextB8)
        m_staticText123 = wx.StaticText(self, wx.ID_ANY, u"额定频率Hz", pos=(THIRD_COL_LABEL, LABEL_TEXT_HEIGHT),
                                        size=(80, 20))
       # m_staticText123.SetFont(wx.Font(9, wx.DEFAULT, wx.NORMAL, wx.NORMAL))
        selectHzs = ['50', '60']
        self.comboxC8 = wx.ComboBox(self, wx.ID_ANY, choices=selectHzs, pos=(THIRD_COL_TEXT, LABEL_TEXT_HEIGHT),
                                    size=(TEXT_WIDTH, 20))
        self.comboxC8.Select(0)
        self.comboxC8.Bind(wx.EVT_COMBOBOX, self.SelectComboxC8)
        m_staticText124 = wx.StaticText(self, wx.ID_ANY, u"相数", pos=(FOURTH_COL_LABEL, LABEL_TEXT_HEIGHT), size=(30, 20))
        selectXiangShu = ['3', '1']
        self.comboxD8 = wx.ComboBox(self, wx.ID_ANY, choices=selectXiangShu, pos=(FOURTH_COL_TEXT, LABEL_TEXT_HEIGHT),
                                    size=(TEXT_WIDTH, 20))
        self.comboxD8.Select(0)
        self.comboxD8.Bind(wx.EVT_COMBOBOX, self.SelectComboxD8)
        # 第-。1行
        m_staticText125 = wx.StaticText(self, wx.ID_ANY, u"冷却方式", pos=(FIRST_COL_LABEL, LABEL_TEXT_HEIGHT+55),
                                        size=(60, 20))
       # m_staticText125.SetFont(wx.Font(9, wx.DEFAULT, wx.NORMAL, wx.NORMAL))
        coolStyle = ["KFWF"]
        self.comboxE8 = wx.ComboBox(self, wx.ID_ANY, choices=coolStyle, pos=(FIRST_COL_TEXT, LABEL_TEXT_HEIGHT+55),
                                    size=(TEXT_WIDTH, 20))
        self.comboxE8.Select(0)
        self.comboxE8.Bind(wx.EVT_COMBOBOX, self.SelectComboxE8)
        m_staticText127 = wx.StaticText(self, wx.ID_ANY, u"海拔高度m", pos=(SECOND_COL_LABEL, LABEL_TEXT_HEIGHT+55),
                                        size=(60, 20))
        self.m_textCtrlK8 = wx.TextCtrl(self, wx.ID_ANY, u"1000", pos=(SECOND_COL_TEXT, LABEL_TEXT_HEIGHT+55), size=(TEXT_WIDTH, 20),
                                        style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlK8.Bind(wx.EVT_TEXT_ENTER, self.EndTextK8)
        m_staticText35 = wx.StaticText(self, wx.ID_ANY, u"基准温度℃", pos=(THIRD_COL_LABEL, LABEL_TEXT_HEIGHT+55),
                                       size=(70, 20))
        self.m_textCtrlO17 = wx.TextCtrl(self, wx.ID_ANY, u"120", pos=(THIRD_COL_TEXT, LABEL_TEXT_HEIGHT+55), size=(TEXT_WIDTH, 20),
                                         style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlO17.Bind(wx.EVT_TEXT_ENTER, self.EndTextO17)
        m_staticText47 = wx.StaticText(self, wx.ID_ANY, u"预选磁密T", pos=(FOURTH_COL_LABEL, LABEL_TEXT_HEIGHT+55),
                                       size=(60, 20))
        self.m_textCtrlB28 = wx.TextCtrl(self, wx.ID_ANY, u"1.65", pos=(FOURTH_COL_TEXT, LABEL_TEXT_HEIGHT+55), size=(TEXT_WIDTH, 20),
                                         style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlB28.Bind(wx.EVT_TEXT_ENTER, self.EndTextB28)
        line = wx.StaticLine(self, pos=(0, 115), size=(LINE_WIDTH, LINE_HEIGHT))
        # 第二行
        highV = wx.StaticText(self, label='高\n\n压', pos=(10, 130), size=(20, 40))
        highV.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        highV.SetForegroundColour('red')
        m_staticText12 = wx.StaticText(self, wx.ID_ANY, u"额定电压kV", pos=(FIRST_COL_LABEL, LABEL_TEXT_HEIGHT+110),
                                       size=(70, 30))
       # m_staticText12.SetFont(wx.Font(8, wx.DEFAULT, wx.NORMAL, wx.NORMAL))
        self.m_textCtrlD10 = wx.TextCtrl(self, wx.ID_ANY, u"36.75", pos=(FIRST_COL_TEXT, LABEL_TEXT_HEIGHT+110), size=(TEXT_WIDTH, 20),
                                         style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlD10.Bind(wx.EVT_TEXT_ENTER, self.EndTextD10)
        m_staticText11 = wx.StaticText(self, wx.ID_ANY, u"容量比", pos=(SECOND_COL_LABEL, LABEL_TEXT_HEIGHT+110),
                                       size=(LABEL_WIDTH, 20))
        self.m_textCtrlC10 = wx.TextCtrl(self, wx.ID_ANY, u"100%", pos=(SECOND_COL_TEXT, LABEL_TEXT_HEIGHT+110), size=(TEXT_WIDTH, 20),
                                         style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlC10.Bind(wx.EVT_TEXT_ENTER, self.EndTextC10)
        m_staticText19 = wx.StaticText(self, wx.ID_ANY, u"连接组别", pos=(THIRD_COL_LABEL, LABEL_TEXT_HEIGHT+110),
                                       size=(60, 20))
        k10s = ["Y", "D"]
        self.comboxK10 = wx.ComboBox(self, wx.ID_ANY, choices=k10s, pos=(THIRD_COL_TEXT, LABEL_TEXT_HEIGHT+110),
                                     size=(TEXT_WIDTH, 20))
        self.comboxK10.Select(1)
        self.comboxK10.Bind(wx.EVT_COMBOBOX, self.EndTextK10)
        line1 = wx.StaticLine(self, pos=(0, 170), size=(LINE_WIDTH, LINE_HEIGHT))
        # 第三行
        lowV = wx.StaticText(self, label='低\n\n压', pos=(10, 185), size=(20, 40))
        lowV.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        lowV.SetForegroundColour('red')
        m_staticText26 = wx.StaticText(self, wx.ID_ANY, u"额定电压kV", pos=(FIRST_COL_LABEL, LABEL_TEXT_HEIGHT+165),
                                       size=(70, 30))
       # m_staticText26.SetFont(wx.Font(8, wx.DEFAULT, wx.NORMAL, wx.NORMAL))
        self.m_textCtrlD14 = wx.TextCtrl(self, wx.ID_ANY, u"0.69", pos=(FIRST_COL_TEXT, LABEL_TEXT_HEIGHT+165), size=(TEXT_WIDTH, 20),
                                         style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlD14.Bind(wx.EVT_TEXT_ENTER, self.EndTextD14)
        m_staticText25 = wx.StaticText(self, wx.ID_ANY, u"容量比", pos=(SECOND_COL_LABEL, LABEL_TEXT_HEIGHT+165),
                                       size=(60, 20))
        self.m_textCtrlC14 = wx.TextCtrl(self, wx.ID_ANY, u"100%", pos=(SECOND_COL_TEXT, LABEL_TEXT_HEIGHT+165), size=(TEXT_WIDTH, 20),
                                         style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlC14.Bind(wx.EVT_TEXT_ENTER, self.EndTextC14)
        m_staticText27 = wx.StaticText(self, wx.ID_ANY, u"连接组别", pos=(THIRD_COL_LABEL, LABEL_TEXT_HEIGHT+165),
                                       size=(60, 20))
        k10s = ["Y", "D"]
        self.comBoxK14 = wx.ComboBox(self, wx.ID_ANY, choices=k10s, pos=(THIRD_COL_TEXT, LABEL_TEXT_HEIGHT+165),
                                     size=(TEXT_WIDTH, 20))
        self.comBoxK14.Select(0)
        self.comBoxK14.Bind(wx.EVT_COMBOBOX, self.EndTextK14)
        line2 = wx.StaticLine(self, pos=(0, 225), size=(LINE_WIDTH, LINE_HEIGHT))

        # 第4行
        adaptVParams = wx.StaticText(self, label='调\n\n压\n\n参\n\n数', pos=(10, 240), size=(20, 80))
        adaptVParams.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        adaptVParams.SetForegroundColour('red')
        m_staticText13 = wx.StaticText(self, wx.ID_ANY, u"调压分接+", pos=(FIRST_COL_LABEL, LABEL_TEXT_HEIGHT+220),
                                       size=(60, 20))
        self.m_textCtrlF10 = wx.TextCtrl(self, wx.ID_ANY, u"2", pos=(FIRST_COL_TEXT, LABEL_TEXT_HEIGHT+220),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlF10.Bind(wx.EVT_TEXT_ENTER, self.EndTextF10)
        m_staticText14 = wx.StaticText(self, wx.ID_ANY, u"调压分接-", pos=(SECOND_COL_LABEL, LABEL_TEXT_HEIGHT+220),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlF11 = wx.TextCtrl(self, wx.ID_ANY, u"2", pos=(SECOND_COL_TEXT, LABEL_TEXT_HEIGHT+220),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlF11.Bind(wx.EVT_TEXT_ENTER, self.EndTextF11)
        m_staticText15 = wx.StaticText(self, wx.ID_ANY, u"级电压比%+", pos=(THIRD_COL_LABEL, LABEL_TEXT_HEIGHT+220),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlG10 = wx.TextCtrl(self, wx.ID_ANY, u"2.5%", pos=(THIRD_COL_TEXT, LABEL_TEXT_HEIGHT+220),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlG10.Bind(wx.EVT_TEXT_ENTER, self.EndTextG10)
        m_staticText16 = wx.StaticText(self, wx.ID_ANY, u"级电压比%-", pos=(FOURTH_COL_LABEL, LABEL_TEXT_HEIGHT+220),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlG11 = wx.TextCtrl(self, wx.ID_ANY, u"2.5%", pos=(FOURTH_COL_TEXT, LABEL_TEXT_HEIGHT+220),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlG11.Bind(wx.EVT_TEXT_ENTER, self.EndTextG11)

        m_staticTextH10 = wx.StaticText(self, wx.ID_ANY, u"调压方式", pos=(FIRST_COL_LABEL, LABEL_TEXT_HEIGHT+275),
                                        size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        H10s = ["无载", "有载"]
        self.comBoxH10 = wx.ComboBox(self, wx.ID_ANY, choices=H10s, pos=(FIRST_COL_TEXT, LABEL_TEXT_HEIGHT+275),
                                     size=(TEXT_WIDTH, 20))
        self.comBoxH10.Bind(wx.EVT_COMBOBOX, self.SelectComH10)
        self.comBoxH10.Select(0)
        m_staticTextI10 = wx.StaticText(self, wx.ID_ANY, u"调压位置", pos=(SECOND_COL_LABEL, LABEL_TEXT_HEIGHT+275),
                                        size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        I10Choices = [u"调压线圈", u"调压段"]
        self.comboxI10 = wx.ComboBox(self, wx.ID_ANY, pos=(SECOND_COL_TEXT, LABEL_TEXT_HEIGHT+275), size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT),
                                     choices=I10Choices)
        self.comboxI10.Bind(wx.EVT_COMBOBOX, self.SelectComI10)
        self.comboxI10.Select(1)
        m_staticTextJ10 = wx.StaticText(self, wx.ID_ANY, u"调压型式", pos=(THIRD_COL_LABEL, LABEL_TEXT_HEIGHT+275),
                                        size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        BoxJ10Choices = [u"正反", u"线性"]
        self.comBoxJ10 = wx.ComboBox(self, wx.ID_ANY, pos=(THIRD_COL_TEXT, LABEL_TEXT_HEIGHT+275), size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT),
                                     choices=BoxJ10Choices)
        self.comBoxJ10.Select(1)
        self.comBoxJ10.Bind(wx.EVT_COMBOBOX, self.SelectComJ10)
        line3 = wx.StaticLine(self, pos=(0, 335), size=(LINE_WIDTH, LINE_HEIGHT))
        # 第5行
        protolParams = wx.StaticText(self, label='协\n\n议\n\n参\n\n数', pos=(10, LABEL_TEXT_HEIGHT+330), size=(20, 80))
        protolParams.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        protolParams.SetForegroundColour('red')
        m_staticText36 = wx.StaticText(self, wx.ID_ANY, u"空载损耗kW", pos=(FIRST_COL_LABEL, LABEL_TEXT_HEIGHT+330),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlC20 = wx.TextCtrl(self, wx.ID_ANY, u"5.5", pos=(FIRST_COL_TEXT, LABEL_TEXT_HEIGHT+330),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlC20.Bind(wx.EVT_TEXT_ENTER, self.EndTextC20)

        m_staticText37 = wx.StaticText(self, wx.ID_ANY, u"空载损耗允差%", pos=(SECOND_COL_LABEL, LABEL_TEXT_HEIGHT+330),
                                       size=(LABEL_WIDTH + 10, LABEL_TEXT_HEIGHT))
        m_staticText37.SetFont(wx.Font(9, wx.DEFAULT, wx.NORMAL, wx.NORMAL))
        self.m_textCtrlC23 = wx.TextCtrl(self, wx.ID_ANY, u"6%", pos=(SECOND_COL_TEXT + 10, LABEL_TEXT_HEIGHT+330),
                                         size=(TEXT_WIDTH - 10, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlC23.Bind(wx.EVT_TEXT_ENTER, self.EndTextC23)
        m_staticText38 = wx.StaticText(self, wx.ID_ANY, u"负载损耗kW", pos=(THIRD_COL_LABEL, LABEL_TEXT_HEIGHT+330),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlE20 = wx.TextCtrl(self, wx.ID_ANY, u"82", pos=(THIRD_COL_TEXT, LABEL_TEXT_HEIGHT+330),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlE20.Bind(wx.EVT_TEXT_ENTER, self.EndTextE20)

        m_staticText39 = wx.StaticText(self, wx.ID_ANY, u"负载损耗允差%", pos=(FOURTH_COL_LABEL, LABEL_TEXT_HEIGHT+330),
                                       size=(LABEL_WIDTH + 20, LABEL_TEXT_HEIGHT))
        self.m_textCtrlE23 = wx.TextCtrl(self, wx.ID_ANY, u"4%", pos=(FOURTH_COL_TEXT, LABEL_TEXT_HEIGHT+330),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlE23.Bind(wx.EVT_TEXT_ENTER, self.EndTextE23)
        m_staticText40 = wx.StaticText(self, wx.ID_ANY, u"阻抗%", pos=(FIRST_COL_LABEL, LABEL_TEXT_HEIGHT+385),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlH20 = wx.TextCtrl(self, wx.ID_ANY, u"9%", pos=(FIRST_COL_TEXT, LABEL_TEXT_HEIGHT+385),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlH20.Bind(wx.EVT_TEXT_ENTER, self.EndTextH20)
        m_staticText41 = wx.StaticText(self, wx.ID_ANY, u"阻抗允差%", pos=(SECOND_COL_LABEL, LABEL_TEXT_HEIGHT+385),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))

        self.m_textCtrlH23 = wx.TextCtrl(self, wx.ID_ANY, u"10%", pos=(SECOND_COL_TEXT, LABEL_TEXT_HEIGHT+385),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlH23.Bind(wx.EVT_TEXT_ENTER, self.EndTextH23)
        m_staticText43 = wx.StaticText(self, wx.ID_ANY, u"噪音dB", pos=(THIRD_COL_LABEL, LABEL_TEXT_HEIGHT+385),
                                       size=(LABEL_WIDTH, 30))
        self.m_textCtrlK20 = wx.TextCtrl(self, wx.ID_ANY, u"68", pos=(THIRD_COL_TEXT, LABEL_TEXT_HEIGHT+385),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlK20.Bind(wx.EVT_TEXT_ENTER, self.EndTextK20)
        m_staticText44 = wx.StaticText(self, wx.ID_ANY, u"噪音允差%", pos=(FOURTH_COL_LABEL, LABEL_TEXT_HEIGHT+385),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlK23 = wx.TextCtrl(self, wx.ID_ANY, u"0%", pos=(FOURTH_COL_TEXT, LABEL_TEXT_HEIGHT+385),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlK23.Bind(wx.EVT_TEXT_ENTER, self.EndTextK23)
        vLine = wx.StaticLine(self, pos=(30, 0), size=(2, 600))

    def __del__(self):
        pass

    def EndTextK23(self, event):
        k23Value = self.m_textCtrlK23.GetValue()
        sheet.range('k23').value = k23Value
        event.Skip()

    def EndTextK20(self, event):
        k20Value = self.m_textCtrlK20.GetValue()
        sheet.range('k20').value = k20Value
        event.Skip()

    def EndTextC3(self, event):
        c3Value = self.m_textCtrlC3.GetValue()
        sheet.range('c3').value = c3Value
        event.Skip()

    def EndTextB8(self, event):
        b8Value = self.m_textCtrlB8.GetValue()
        sheet.range('B8').value = b8Value
        event.Skip()

    def EndTextK8(self, event):
        k8Value = self.m_textCtrlK8.GetValue()
        sheet.range('k8').value = k8Value
        event.Skip()

    def EndTextO17(self, event):
        o17Value = self.m_textCtrlO17.GetValue()
        sheet.range('o17').value = o17Value
        event.Skip()

    def EndTextB28(self, event):
        b28Value = self.m_textCtrlB8.GetValue()
        sheet.range('b28').value = b28Value
        event.Skip()

    def EndTextD10(self, event):
        d10Value = self.m_textCtrlD10.GetValue()
        sheet.range('d10').value = d10Value
        event.Skip()

    def EndTextC10(self, event):
        c10Value = self.m_textCtrlC10.GetValue()
        sheet.range('c10').value = c10Value
        event.Skip()

    def EndTextD14(self, event):
        d14Value = self.m_textCtrlD14.GetValue()
        sheet.range('d14').value = d14Value
        event.Skip()

    def EndTextC14(self, event):
        c14Value = self.m_textCtrlC14.GetValue()
        sheet.range('c14').value = c14Value
        event.Skip()

    def EndTextF10(self, event):
        f10Value = self.m_textCtrlF10.GetValue()
        sheet.range('f10').value = f10Value
        event.Skip()

    def EndTextK10(self, event):
        k10Value = self.comboxK10.GetValue()
        sheet.range('k10').value = k10Value
        event.Skip()

    def EndTextK14(self, event):
        k14Value = self.comBoxK14.GetValue()
        sheet.range('k14').value = k14Value
        event.Skip()

    def SelectComboxC8(self, event):
        c8Value = self.comboxC8.GetValue()
        sheet.range('c8').value = c8Value
        event.Skip()

    def SelectComboxD8(self, event):
        d8Value = self.comboxD8.GetValue()
        sheet.range('d8').value = d8Value
        event.Skip()

    def SelectComboxE8(self, event):
        e8Value = self.comboxE8.Value
        sheet.range('e8').value = e8Value
        event.Skip()

    def SelectComH10(self, event):
        h10Value = self.comBoxH10.Value
        sheet.range('h10').value = h10Value
        event.Skip()

    def SelectComI10(self, event):
        i10Value = self.comboxI10.Value
        sheet.range('i10').value = i10Value
        event.Skip()

    def SelectComJ10(self, event):
        j10Value = self.comBoxJ10.Value
        sheet.range('j10').value = j10Value
        event.Skip()

    def EndTextF11(self, event):
        f11Value = self.m_textCtrlF11.GetValue()
        sheet.range('f11').value = f11Value
        event.Skip()

    def EndTextG10(self, event):
        g10Value = self.m_textCtrlG10.GetValue()
        sheet.range('g10').value = g10Value
        event.Skip()

    def EndTextG11(self, event):
        g11Value = self.m_textCtrlG11.GetValue()
        sheet.range('g11').value = g11Value
        event.Skip()

    def EndTextC20(self, event):
        c20Value = self.m_textCtrlC20.GetValue()
        sheet.range('c20').value = c20Value
        event.Skip()

    def EndTextC23(self, event):
        c23Value = self.m_textCtrlC23.GetValue()
        sheet.range('c23').value = c23Value
        event.Skip()

    def EndTextE20(self, event):
        e20Value = self.m_textCtrlE20.GetValue()
        sheet.range('e20').value = e20Value
        event.Skip()

    def EndTextE23(self, event):
        e23Value = self.m_textCtrlE23.GetValue()
        sheet.range('e23').value = e23Value
        event.Skip()

    def EndTextH20(self, event):
        h20Value = self.m_textCtrlH20.GetValue()
        sheet.range('h20').value = h20Value
        event.Skip()

    def EndTextH23(self, event):
        h23Value = self.m_textCtrlH23.GetValue()
        sheet.range('h23').value = h23Value
        event.Skip()


class MyPanel2(wx.Panel):
    def __init__(self, parent):
        super(MyPanel2, self).__init__(parent)
        vLine = wx.StaticLine(self, pos=(30, 0), size=(2, 600))
        # 第 0 行 线圈排列
        circleLine = wx.StaticText(self, label='线\n\n圈\n\n排\n\n列', pos=(FIRST_ELE_X, FIRST_ELE_Y), size=(20, 80))
        circleLine.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        circleLine.SetForegroundColour('red')
        B25Choices = [u"HV", u"LV"]

        m_staticText45 = wx.StaticText(self, wx.ID_ANY, u"线圈排列1", pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 20),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.comboxB25 = wx.ComboBox(self, wx.ID_ANY, pos=(FIRST_COL_TEXT, FIRST_ELE_Y + 20),
                                     size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), choices=B25Choices)
        self.comboxB25.Select(1)
        self.comboxB25.Bind(wx.EVT_COMBOBOX, self.SelectComboxB25)
        m_staticText46 = wx.StaticText(self, wx.ID_ANY, u"线圈排列2", pos=(SECOND_COL_LABEL, FIRST_ELE_Y + 20),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.comboxC25 = wx.ComboBox(self, wx.ID_ANY, pos=(SECOND_COL_TEXT, FIRST_ELE_Y + 20),
                                     size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), choices=B25Choices)
        self.comboxC25.Select(0)
        self.comboxC25.Bind(wx.EVT_COMBOBOX, self.SelectComboxC25)
        line0 = wx.StaticLine(self, pos=(0, 100), size=(LINE_WIDTH, LINE_HEIGHT))
        # 第一行
        lowV = wx.StaticText(self, label='低\n\n压', pos=(FIRST_ELE_X, FIRST_ELE_Y + 115),
                             size=(LABEL_TEXT_HEIGHT, 60))
        lowV.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        lowV.SetForegroundColour('red')
        m_radioBoxB38Choices = [u"连续式", u"螺旋式", u"箔绕式"]
        m_staticText38 = wx.StaticText(self, wx.ID_ANY, u"线圈型式", pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 100),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.comboxB38 = wx.ComboBox(self, wx.ID_ANY, pos=(FIRST_COL_TEXT, FIRST_ELE_Y + 100),
                                     size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), choices=m_radioBoxB38Choices)
        self.comboxB38.Select(2)
        self.comboxB38.Bind(wx.EVT_COMBOBOX, self.SelectComboxB38)
        m_staticTextG38 = wx.StaticText(self, wx.ID_ANY, u"层数1", pos=(SECOND_COL_LABEL, FIRST_ELE_Y + 100),
                                        size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.textCtrlG38 = wx.TextCtrl(self, wx.ID_ANY, u"2", pos=(SECOND_COL_TEXT, FIRST_ELE_Y + 100),
                                       size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.textCtrlG38.Bind(wx.EVT_TEXT_ENTER, self.EndTextG38)

        m_staticText50 = wx.StaticText(self, wx.ID_ANY, u"层数2", pos=(THIRD_COL_LABEL, FIRST_ELE_Y + 100),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlH38 = wx.TextCtrl(self, wx.ID_ANY, u"2", pos=(THIRD_COL_TEXT, FIRST_ELE_Y + 100),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlH38.Bind(wx.EVT_TEXT_ENTER, self.EndTextH38)
        m_staticText51 = wx.StaticText(self, wx.ID_ANY, u"并联1", pos=(FOURTH_COL_LABEL, FIRST_ELE_Y + 100),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlG39 = wx.TextCtrl(self, wx.ID_ANY, u"3", pos=(FOURTH_COL_TEXT, FIRST_ELE_Y + 100),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlG39.Bind(wx.EVT_TEXT_ENTER, self.EndTextG39)
        m_staticText53 = wx.StaticText(self, wx.ID_ANY, u"并联2", pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 130),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlH39 = wx.TextCtrl(self, wx.ID_ANY, u"2", pos=(FIRST_COL_TEXT, FIRST_ELE_Y + 130),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlH39.Bind(wx.EVT_TEXT_ENTER, self.EndTextH39)
        m_staticText55 = wx.StaticText(self, wx.ID_ANY, u"大线并联根数", pos=(SECOND_COL_LABEL, FIRST_ELE_Y + 130),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlF40 = wx.TextCtrl(self, wx.ID_ANY, u"1", pos=(SECOND_COL_TEXT, FIRST_ELE_Y + 130),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlF40.Bind(wx.EVT_TEXT_ENTER, self.EndTextF40)

        B43BoxChoices = [u"普通", u"复合", u"换位", u"箔式"]
        m_staticTextB43 = wx.StaticText(self, wx.ID_ANY, u"导线型式", pos=(THIRD_COL_LABEL, FIRST_ELE_Y + 130),
                                        size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.BoxB43 = wx.ComboBox(self, wx.ID_ANY, pos=(THIRD_COL_TEXT, FIRST_ELE_Y + 130),
                                  size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), choices=B43BoxChoices)
        self.BoxB43.Select(3)
        self.BoxB43.Bind(wx.EVT_COMBOBOX, self.SelectComboxB43)
        m_staticText57 = wx.StaticText(self, wx.ID_ANY, u"裸线尺寸厚mm", pos=(FOURTH_COL_LABEL, FIRST_ELE_Y + 130),
                                       size=(LABEL_WIDTH + 20, LABEL_TEXT_HEIGHT))
        self.m_textCtrlB44 = wx.TextCtrl(self, wx.ID_ANY, u"1.4", pos=(FOURTH_COL_TEXT, FIRST_ELE_Y + 130),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlB44.Bind(wx.EVT_TEXT_ENTER, self.EndTextB44)

        m_staticText59 = wx.StaticText(self, wx.ID_ANY, u"裸线尺寸宽mm", pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 160),
                                       size=(LABEL_WIDTH + 20, LABEL_TEXT_HEIGHT))
        self.m_textCtrlD44 = wx.TextCtrl(self, wx.ID_ANY, u"920", pos=(FIRST_COL_TEXT + 20, FIRST_ELE_Y + 160),
                                         size=(TEXT_WIDTH - 20, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlD44.Bind(wx.EVT_TEXT_ENTER, self.EndTextD44)

        sm_staticText60 = wx.StaticText(self, wx.ID_ANY, u"每并联根数", pos=(SECOND_COL_LABEL, FIRST_ELE_Y + 160),
                                        size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlG45 = wx.TextCtrl(self, wx.ID_ANY, u"2", pos=(SECOND_COL_TEXT, FIRST_ELE_Y + 160),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlG45.Bind(wx.EVT_TEXT_ENTER, self.EndTextG45)

        m_staticText62 = wx.StaticText(self, wx.ID_ANY, u"档油圈数", pos=(THIRD_COL_LABEL, FIRST_ELE_Y + 160),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlH47 = wx.TextCtrl(self, wx.ID_ANY, u"0", pos=(THIRD_COL_TEXT, FIRST_ELE_Y + 160),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlH47.Bind(wx.EVT_TEXT_ENTER, self.EndTextH47)

        m_staticText63 = wx.StaticText(self, wx.ID_ANY, u"油道高mm", pos=(FOURTH_COL_LABEL, FIRST_ELE_Y + 160),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlD56 = wx.TextCtrl(self, wx.ID_ANY, u"4", pos=(FOURTH_COL_TEXT, FIRST_ELE_Y + 160),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlD56.Bind(wx.EVT_TEXT_ENTER, self.EndTextD56)

        line1 = wx.StaticLine(self, pos=(0, 195), size=(LINE_WIDTH, LINE_HEIGHT))
        # 第2行
        heightV = wx.StaticText(self, label='高\n\n压', pos=(FIRST_ELE_X, FIRST_ELE_Y + 215),
                                size=(LABEL_TEXT_HEIGHT, 50))
        heightV.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        heightV.SetForegroundColour('red')
        m_radioBox8B71Choices = [u"连续式", u"纠连式"]
        m_staticTextB71 = wx.StaticText(self, wx.ID_ANY, u"线圈型式", pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 200),
                                        size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.BoxB71 = wx.ComboBox(self, wx.ID_ANY, pos=(FIRST_COL_TEXT, FIRST_ELE_Y + 200),
                                  size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), choices=m_radioBox8B71Choices)
        self.BoxB71.Select(1)
        self.BoxB71.Bind(wx.EVT_COMBOBOX, self.SelectComboxB71)

        m_staticText72 = wx.StaticText(self, wx.ID_ANY, u"线段数", pos=(SECOND_COL_LABEL, FIRST_ELE_Y + 200),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlB73 = wx.TextCtrl(self, wx.ID_ANY, u"84", pos=(SECOND_COL_TEXT, FIRST_ELE_Y + 200),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlB73.Bind(wx.EVT_TEXT_ENTER, self.EndTextB73)

        B76BoxChoices = [u"普通", u"复合", u"换位"]
        m_staticTextB76 = wx.StaticText(self, wx.ID_ANY, u"导线型式", pos=(THIRD_COL_LABEL, FIRST_ELE_Y + 200),
                                        size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.comBoxB76 = wx.ComboBox(self, wx.ID_ANY, pos=(THIRD_COL_TEXT, FIRST_ELE_Y + 200),
                                     size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), choices=B76BoxChoices)
        self.comBoxB76.Bind(wx.EVT_COMBOBOX, self.SelectComboxB76)
        self.comBoxB76.Select(1)

        m_staticTextB77 = wx.StaticText(self, wx.ID_ANY, u"裸线尺寸厚mm", pos=(FOURTH_COL_LABEL, FIRST_ELE_Y + 200),
                                        size=(LABEL_WIDTH + 20, LABEL_TEXT_HEIGHT))
        self.m_textCtrlB77 = wx.TextCtrl(self, wx.ID_ANY, u"2.84", pos=(FOURTH_COL_TEXT, FIRST_ELE_Y + 200),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlB77.Bind(wx.EVT_TEXT_ENTER, self.EndTextB77)

        m_staticTextD77 = wx.StaticText(self, wx.ID_ANY, u"裸线尺寸宽mm", pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 240),
                                        size=(LABEL_WIDTH + 20, LABEL_TEXT_HEIGHT))
        self.m_textCtrlD77 = wx.TextCtrl(self, wx.ID_ANY, u"7.4", pos=(FIRST_COL_TEXT + 20, FIRST_ELE_Y + 240),
                                         size=(TEXT_WIDTH - 20, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlD77.Bind(wx.EVT_TEXT_ENTER, self.EndTextD77)

        sm_staticG78 = wx.StaticText(self, wx.ID_ANY, u"并联根数", pos=(SECOND_COL_LABEL, FIRST_ELE_Y + 240),
                                     size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlG78 = wx.TextCtrl(self, wx.ID_ANY, u"2", pos=(SECOND_COL_TEXT, FIRST_ELE_Y + 240),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlG78.Bind(wx.EVT_TEXT_ENTER, self.EndTextG78)

        m_staticH80 = wx.StaticText(self, wx.ID_ANY, u"档油圈数", pos=(THIRD_COL_LABEL, FIRST_ELE_Y + 240),
                                    size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlH80 = wx.TextCtrl(self, wx.ID_ANY, u"0", pos=(THIRD_COL_TEXT, FIRST_ELE_Y + 240),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlH80.Bind(wx.EVT_TEXT_ENTER, self.EndTextH80)

        m_staticD89 = wx.StaticText(self, wx.ID_ANY, u"油道高mm", pos=(FOURTH_COL_LABEL, FIRST_ELE_Y + 240),
                                    size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlD89 = wx.TextCtrl(self, wx.ID_ANY, u"4", pos=(FOURTH_COL_TEXT, FIRST_ELE_Y + 240),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlD89.Bind(wx.EVT_TEXT_ENTER, self.EndTextD89)

        line2 = wx.StaticLine(self, pos=(0, 285), size=(LINE_WIDTH, LINE_HEIGHT))
        # 第3行 线圈绝缘半径
        row = wx.StaticText(self, label='线圈\n\n绝缘\n\n半径', pos=(0, FIRST_ELE_Y + 300),
                            size=(LABEL_TEXT_HEIGHT, 80))
        row.SetFont(wx.Font(wx.DEFAULT, wx.DEFAULT, wx.DEFAULT, wx.FONTWEIGHT_BOLD))
        row.SetForegroundColour('red')
        m_staticText31 = wx.StaticText(self, wx.ID_ANY, u"铁心直径", pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 300),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlB17 = wx.TextCtrl(self, wx.ID_ANY, u"360", pos=(FIRST_COL_TEXT, FIRST_ELE_Y + 300),
                                         size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlB17.Bind(wx.EVT_TEXT_ENTER, self.EndTextB17)

        m_staticText83 = wx.StaticText(self, wx.ID_ANY, u"铁心-低压mm",
                                       pos=(SECOND_COL_LABEL, FIRST_ELE_Y + 300),
                                       size=(LABEL_WIDTH + 20, LABEL_TEXT_HEIGHT))
        self.m_textCtrlB191 = wx.TextCtrl(self, wx.ID_ANY, u"10", pos=(SECOND_COL_TEXT + 20, FIRST_ELE_Y + 300),
                                          size=(TEXT_WIDTH - 20, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlB191.Bind(wx.EVT_TEXT_ENTER, self.EndTextB191)

        m_staticText84 = wx.StaticText(self, wx.ID_ANY, u"低压-高压mm", pos=(THIRD_COL_LABEL, FIRST_ELE_Y + 300),
                                       size=(LABEL_WIDTH + 20, 30))
        self.m_textCtrlB195 = wx.TextCtrl(self, wx.ID_ANY, u"22", pos=(THIRD_COL_TEXT, FIRST_ELE_Y + 300),
                                          size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlB195.Bind(wx.EVT_TEXT_ENTER, self.EndTextB195)

        m_staticText85 = wx.StaticText(self, wx.ID_ANY, u"上部mm", pos=(FOURTH_COL_LABEL, FIRST_ELE_Y + 300),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlI192 = wx.TextCtrl(self, wx.ID_ANY, u"65", pos=(FOURTH_COL_TEXT, FIRST_ELE_Y + 300),
                                          size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlI192.Bind(wx.EVT_TEXT_ENTER, self.EndTextI192)

        m_staticText86 = wx.StaticText(self, wx.ID_ANY, u"下部mm", pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 340),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlI193 = wx.TextCtrl(self, wx.ID_ANY, u"65", pos=(FIRST_COL_TEXT, FIRST_ELE_Y + 340),
                                          size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlI193.Bind(wx.EVT_TEXT_ENTER, self.EndTextI193)

        m_staticText87 = wx.StaticText(self, wx.ID_ANY, u"间隙mm", pos=(SECOND_COL_LABEL, FIRST_ELE_Y + 340),
                                       size=(LABEL_WIDTH, LABEL_TEXT_HEIGHT))
        self.m_textCtrlI194 = wx.TextCtrl(self, wx.ID_ANY, u"10", pos=(SECOND_COL_TEXT, FIRST_ELE_Y + 340),
                                          size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT), style=wx.TE_PROCESS_ENTER)
        self.m_textCtrlI194.Bind(wx.EVT_TEXT_ENTER, self.EndTextI194)

        line3 = wx.StaticLine(self, pos=(0, 390), size=(LINE_WIDTH, LINE_HEIGHT))
        exPortBtn = wx.Button(self, wx.ID_ANY, '导出word结果', pos=(FIRST_COL_LABEL, FIRST_ELE_Y + 390),
                              size=(TEXT_WIDTH, LABEL_TEXT_HEIGHT))
        exPortBtn.Bind(wx.EVT_BUTTON, ExportDoc)

    def __del__(self):
        appEx.quit()
        pass

    def EndTextI193(self, event):
        i195Value = self.m_textCtrlI193.GetValue()
        sheet.range('i193').value = i195Value
        event.Skip()

    def EndTextI194(self, event):
        i194Value = self.m_textCtrlI194.GetValue()
        sheet.range('i194').value = i194Value
        event.Skip()

    def EndTextI192(self, event):
        i192Value = self.m_textCtrlI192.GetValue()
        sheet.range('i192').value = i192Value
        event.Skip()

    def EndTextB195(self, event):
        b195Value = self.m_textCtrlB195.GetValue()
        sheet.range('b195').value = b195Value
        event.Skip()

    def EndTextB191(self, event):
        b191Value = self.m_textCtrlB191.GetValue()
        sheet.range('b191').value = b191Value
        event.Skip()

    def EndTextB17(self, event):
        b17Value = self.m_textCtrlB17.GetValue()
        sheet.range('b17').value = b17Value
        event.Skip()

    def EndTextD89(self, event):
        d89Value = self.m_textCtrlD89.GetValue()
        sheet.range('d89').value = d89Value
        event.Skip()

    def EndTextH80(self, event):
        h80Value = self.m_textCtrlH80.GetValue()
        sheet.range('h80').value = h80Value
        event.Skip()

    def EndTextG78(self, event):
        g78Value = self.m_textCtrlG78.GetValue()
        sheet.range('g78').value = g78Value
        event.Skip()

    def EndTextD77(self, event):
        d77Value = self.m_textCtrlD77.GetValue()
        sheet.range('d77').value = d77Value
        event.Skip()

    def EndTextB77(self, event):
        b77Value = self.m_textCtrlB77.GetValue()
        sheet.range('b77').value = b77Value
        event.Skip()

    def EndTextB73(self, event):
        b73Value = self.m_textCtrlB73.GetValue()
        sheet.range('b73').value = b73Value
        event.Skip()

    def EndTextD56(self, event):
        d56Value = self.m_textCtrlD56.GetValue()
        sheet.range('d56').value = d56Value
        event.Skip()

    def EndTextH47(self, event):
        h47Value = self.m_textCtrlH47.GetValue()
        sheet.range('h47').value = h47Value
        event.Skip()

    def EndTextG45(self, event):
        g45Value = self.m_textCtrlG45.GetValue()
        sheet.range('g45').value = g45Value
        event.Skip()

    def EndTextD44(self, event):
        d44Value = self.m_textCtrlD44.GetValue()
        sheet.range('d44').value = d44Value
        event.Skip()

    def EndTextB44(self, event):
        b44Value = self.m_textCtrlB44.GetValue()
        sheet.range('b44').value = b44Value
        event.Skip()

    def EndTextF40(self, event):
        f40Value = self.m_textCtrlF40.GetValue()
        sheet.range('f40').value = f40Value
        event.Skip()

    def EndTextH39(self, event):
        h39Value = self.m_textCtrlH39.GetValue()
        sheet.range('h39').value = h39Value
        event.Skip()

    def EndTextG39(self, event):
        g39Value = self.m_textCtrlG39.GetValue()
        sheet.range('g39').value = g39Value
        event.Skip()

    def EndTextH38(self, event):
        h38Value = self.m_textCtrlH38.GetValue()
        sheet.range('h38').value = h38Value
        event.Skip()

    def EndTextG38(self, event):
        g38Value = self.m_textCtrlG38.GetValue()
        sheet.range('g38').value = g38Value
        event.Skip()

    def SelectComboxB25(self, event):
        b25Value = self.comboxB25.GetValue()
        sheet.range('b25').value = b25Value
        event.Skip()

    def SelectComboxC25(self, event):
        c25Value = self.comboxC25.GetValue()
        sheet.range('c25').value = c25Value
        event.Skip()

    def SelectComboxB38(self, event):
        b38Value = self.comboxB38.Value
        sheet.range('b38').value = b38Value
        event.Skip()

    def SelectComboxB43(self, event):
        b38Value = self.BoxB43.Value
        sheet.range('b43').value = b38Value
        event.Skip()

    def SelectComboxB71(self, event):
        b71Value = self.BoxB71.Value
        sheet.range('b71').value = b71Value
        event.Skip()

    def SelectComboxB76(self, event):
        b76Value = self.comBoxB76.Value
        sheet.range('b76').value = b76Value
        event.Skip()

